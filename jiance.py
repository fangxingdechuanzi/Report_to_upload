import os
from docx import Document
import re
import openpyxl

class JianCe:
    def __init__(self,shen_he,bian_zhi):
        self.pizhun = '秦专专'
        self.shenhe = shen_he
        self.bianzhi = bian_zhi
        # self.path=path
        self.work_path=os.getcwd()
        self.excel_path=self.work_path+'\\excel\\检验统计.xlsx'
    #打开文件读取信息
    def open_docx(self,path):
        doc=Document(path)
        tables=doc.tables
        paras=doc.paragraphs
        #报告编号
        for para in paras:
            if re.search(r'JC-\d{8,9}',para.text):
                bian_hao=re.search(r'JC-\d{8,9}',para.text)
                b0=bian_hao.group()
                break
        #报告结论表
        jie_lun_biao=tables[2]
        #受检单位
        b1=jie_lun_biao.rows[2].cells[-1].text
        #委托单位
        b2=jie_lun_biao.rows[1].cells[-1].text
        #产品名称
        b3=jie_lun_biao.rows[0].cells[2].text
        #型号规格
        b4=jie_lun_biao.rows[1].cells[2].text
        #检测时间
        a5=jie_lun_biao.rows[4].cells[-1].text
        year = int(a5[:4])
        month = int(a5[5:7])
        day = int(a5[-3:-1])
        b5='{0}/{1}/{2}'.format(year,month,day)
        #检验地点
        b6=jie_lun_biao.rows[5].cells[-1].text
        #检验结论
        jian_yan=jie_lun_biao.rows[8].cells[-1].text
        jian_yan_list=jian_yan.split('\n')
        #下次检验截止日期
        if re.search(r'钻机',b3):
            b7 = '1、' + jian_yan_list[0] + '\n'+jian_yan_list[1]+'\n' + '2、' + jian_yan_list[2] + '\n' + '3、' + jian_yan_list[3]
            ri_qi = re.search(r'\d{4}年\d{1,2}月\d{1,2}日', jian_yan_list[3])
            a8 = ri_qi.group()
        else:
            b7 = '1、' + jian_yan_list[0] + '\n' + '2、' + jian_yan_list[1] + '\n' + '3、' + jian_yan_list[2]
            ri_qi=re.search(r'\d{4}年\d{1,2}月\d{1,2}日',jian_yan_list[2])
            a8=ri_qi.group()
        year1 = int(a8[:4])
        month1 = int(a8[5:7])
        day1 = int(a8[-3:-1])
        b8 = '{0}/{1}/{2}'.format(year1, month1, day1)
        #签发日期
        if month in [1,3,5,7,8,10]:
            if day == 31:
                b9='{0}/{1}/{2}'.format(year,month+1,1)
            else:
                b9='{0}/{1}/{2}'.format(year,month,day+1)
        elif month in [4,6,9,11]:
            if day == 30:
                b9='{0}/{1}/{2}'.format(year,month+1,1)
            else:
                b9='{0}/{1}/{2}'.format(year,month,day+1)
        elif month == 2:
            if day == 28:
                b9='{0}/{1}/{2}'.format(year,month+1,1)
            else:
                b9='{0}/{1}/{2}'.format(year,month,day+1)
        elif month == 12:
            if day == 31:
                b9 = '{0}/{1}/{2}'.format(year+1,1,1)
            else:
                b9='{0}/{1}/{2}'.format(year,month,day+1)

        #批准
        b10=self.pizhun
        #审核
        b11=self.shenhe
        #编制
        b12=self.bianzhi
        self.b_list=[b0, b1, b2, b3, b4, b5, b6, b7, b8, b9, b10, b11, b12]
        return self.b_list

    #复制表格到result目录下
    def copy_excel(self):
        file = self.work_path + '\\result\\井架检验上传统计.xlsx'
        wb=openpyxl.load_workbook(self.excel_path)
        work_sheet = wb['修井机-二维码']
        wb.save(file)

    def open_excel(self):
        self.file = self.work_path + '\\result\\井架检验上传统计.xlsx'
        self.wb = openpyxl.load_workbook(self.file)
        self.work_sheet = self.wb['修井机-二维码']
    #将信息保存到表格
    def write_excel(self,infor,num):
        try:
            for i in range(len(infor)):
                self.work_sheet.cell(row=num+2,column=i+1,value=infor[i])
        except:
            pass

    def save_excel(self):
        self.wb.save(self.file)
