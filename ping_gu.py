import os
from docx import Document
import re
import openpyxl

from first import Change

class PingGu:
    def __init__(self,shen_he,bian_zhi):
        self.pizhun = '刘江涛'
        self.shenhe = shen_he
        self.bianzhi = bian_zhi
        self.work_path=os.getcwd()
        self.excel_path=self.work_path+'\\excel\\评估报告统计上传.xlsx'
    #打开文件读取信息
    def open_docx(self,path):
        doc=Document(path)
        tables=doc.tables
        paras=doc.paragraphs
        #报告编号
        a0=tables[0].rows[0].cells[-1].text
        bian_hao = re.search(r'01202\d{4,5}', a0)
        b0=bian_hao.group()
        # 评估报告结论表
        for table in tables:
            if table.rows[0].cells[0].text=='受评估队伍':
                jie_lun_biao = table
                break
        # 受检单位
        b1=jie_lun_biao.rows[0].cells[2].text
        # 委托单位
        b2=jie_lun_biao.rows[1].cells[2].text
        # 专业类别
        b3=jie_lun_biao.rows[0].cells[-1].text
        # 评估时间
        a4=jie_lun_biao.rows[6].cells[1].text
        if re.search(r'\d{4}\.\d{1,2}\.\d{1,2}',a4):
            a4_time=a4.split('.')
            year = int(a4_time[0])
            month = int(a4_time[1])
            day = int(a4_time[2])
        else:
            year = int(a4[:4])
            month = int(a4[5:7])
            day = int(a4[-3:-1])
        b4='{0}/{1}/{2}'.format(year,month,day)
        #评估地点
        b5=jie_lun_biao.rows[6].cells[-1].text
        #评估结论
        jie_lun=jie_lun_biao.rows[7].cells[-1].text
        a6=jie_lun.split('\n')
        aa6=['1.'+a6[0],a6[1],'2.'+a6[2],'3.'+a6[3]]
        b6='\n'.join(aa6)
        #下次检验截止日期
        #应力测试表
        for table_cs in tables:
            if len(table_cs.rows[0].cells)>2 and re.search(r'井架型号规格',table_cs.rows[0].cells[1].text):
                ji_bie=table_cs.rows[10].cells[-1].text
                a7=re.search(r'\d{4}年\d{1,2}月\d{1,2}日',ji_bie).group()
                year1=int(a7[:4])
                month1=int(a7[5:7])
                day1=int(a7[-3:-1])
                b7='{0}/{1}/{2}'.format(year1,month1,day1)
        #签发日期
        if month in [1,3,5,7,8,10]:
            if day == 31:
                b8='{0}/{1}/{2}'.format(year,month+1,1)
            else:
                b8='{0}/{1}/{2}'.format(year,month,day+1)
        elif month in [4,6,9,11]:
            if day == 30:
                b8='{0}/{1}/{2}'.format(year,month+1,1)
            else:
                b8='{0}/{1}/{2}'.format(year,month,day+1)
        elif month == 2:
            if day == 28:
                b8='{0}/{1}/{2}'.format(year,month+1,1)
            else:
                b8='{0}/{1}/{2}'.format(year,month,day+1)
        elif month == 12:
            if day == 31:
                b8 = '{0}/{1}/{2}'.format(year+1,1,1)
            else:
                b8='{0}/{1}/{2}'.format(year,month,day+1)
        #批准
        b9=self.pizhun
        #审核
        b10=self.shenhe
        #编制
        b11=self.bianzhi
        self.b_list=[b0, b1, b2, b3, b4, b5, b6, b7, b8, b9, b10, b11]
        return self.b_list

    #复制表格到result目录下
    def copy_excel(self):
        file = self.work_path + '\\result\\评估上传统计.xlsx'
        wb=openpyxl.load_workbook(self.excel_path)
        work_sheet = wb['Sheet1']
        wb.save(file)

    def open_excel(self):
        self.file = self.work_path + '\\result\\评估上传统计.xlsx'
        self.wb = openpyxl.load_workbook(self.file)
        self.work_sheet = self.wb['Sheet1']
    #将信息保存到表格
    def write_excel(self,infor,num):
        try:
            for i in range(len(infor)):
                self.work_sheet.cell(row=num+2,column=i+1,value=infor[i])
        except:
            pass
        # for i in range(len(infor)):
        #     self.work_sheet.cell(row=num+2,column=i+1,value=infor[i])

    def save_excel(self):
        self.wb.save(self.file)

#-----------------测试-----------------
# ping_gu_bao_gao=Change().output3()
# ax=PingGu('丁欢欢','王乡桷')
# nn=ax.open_docx(ping_gu_bao_gao[0])
# ax.copy_excel()
# ax.open_excel()
# ax.write_excel(nn,0)
# ax.save_excel()
