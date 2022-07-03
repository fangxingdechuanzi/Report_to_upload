import os
from docx import Document
import re
import openpyxl
class TanShang:
    def __init__(self,bian_zhi):
        self.pizhun = '方义平'
        self.bianzhi = bian_zhi
        self.work_path=os.getcwd()
        self.excel_path=self.work_path+'\\excel\\探伤检测统计.xlsx'
    #打开文件读取信息
    def open_docx(self,path):
        print(path)
        doc=Document(path)
        print(doc)
        tables=doc.tables
        paras=doc.paragraphs
        #报告编号
        for para in paras:
            if re.search(r'JC-\d{8,9}',para.text):
                bian_hao=re.search(r'JC-\d{8,9}',para.text)
                b0=bian_hao.group()
                break
        #报告结论表
        jie_lun_biao=tables[2]
        #受检单位
        b1=jie_lun_biao.rows[1].cells[-1].text
        #委托单位
        b2=jie_lun_biao.rows[1].cells[2].text
        #产品名称
        b3=jie_lun_biao.rows[0].cells[2].text
        #检测项目
        b4='磁粉检测'
        #检测时间
        a5=jie_lun_biao.rows[4].cells[-1].text
        year = int(a5[:4])
        month = int(a5[5:7])
        day = int(a5[-3:-1])
        b5='{0}/{1}/{2}'.format(year,month,day)
        #检验地点
        b6=jie_lun_biao.rows[5].cells[-1].text
        #检验结论
        b7=jie_lun_biao.rows[8].cells[-1].text
        #签发日期
        if month in [1,3,5,7,8,10]:
            if day == 31:
                b8='{0}/{1}/{2}'.format(year,month+1,1)
            else:
                b8='{0}/{1}/{2}'.format(year,month,day+1)
        elif month in [4,6,9,11]:
            if day == 30:
                b8='{0}/{1}/{2}'.format(year,month+1,1)
            else:
                b8='{0}/{1}/{2}'.format(year,month,day+1)
        elif month == 2:
            if day == 28:
                b8='{0}/{1}/{2}'.format(year,month+1,1)
            else:
                b8='{0}/{1}/{2}'.format(year,month,day+1)
        elif month == 12:
            if day == 31:
                b8 = '{0}/{1}/{2}'.format(year+1,1,1)
            else:
                b8='{0}/{1}/{2}'.format(year,month,day+1)

        #批准
        b10=self.pizhun
        #编制
        b9=self.bianzhi
        self.b_list=[b0, b1, b2, b3, b4, b5, b6, b7, b8, b9,b10]
        return self.b_list

    #复制表格到result目录下
    def copy_excel(self):
        file = self.work_path + '\\result\\探伤上传统计.xlsx'
        wb=openpyxl.load_workbook(self.excel_path)
        work_sheet = wb['探伤二维码']
        wb.save(file)

    def open_excel(self):
        self.file = self.work_path + '\\result\\探伤上传统计.xlsx'
        self.wb = openpyxl.load_workbook(self.file)
        self.work_sheet = self.wb['探伤二维码']
    #将信息保存到表格
    def write_excel(self,infor,num):
        for i in range(len(infor)):
            self.work_sheet.cell(row=num+2,column=i+1,value=infor[i])

    def save_excel(self):
        self.wb.save(self.file)

# def jian_ce_tan_shang(pi_zhun,bian_zhi):
#     b = Change().output2()
#     tan_shang = TanShang(pi_zhun, bian_zhi)
#     tan_shang.copy_excel()
#     tan_shang.open_excel()
#     for i in range(len(b)):
#         ab = tan_shang.open_docx(b[i])
#         tan_shang.write_excel(ab, i)
#         os.remove(b[i])
#     tan_shang.save_excel()
# jian_ce_tan_shang('马洋洋','方义平')